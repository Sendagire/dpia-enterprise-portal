import streamlit as st
import google.generativeai as genai
import datetime
from docx import Document
from supabase import create_client
from io import BytesIO

# --- 1. CONFIG & STYLING ---
st.set_page_config(page_title="DPIA Enterprise", layout="wide")

st.markdown("""
    <style>
    .stApp { background-color: #f0f2f5; }
    .sub-header { 
        font-size: 13px; font-weight: 800; color: #0866ff; 
        border-bottom: 2px solid #f0f2f5; padding-bottom: 4px; 
        margin-top: 20px; margin-bottom: 12px; text-transform: uppercase; 
    }
    div[data-testid="stForm"] { background: white; padding: 20px; border-radius: 8px; border: 1px solid #e5e7eb; }
    </style>
""", unsafe_allow_html=True)

genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
model = genai.GenerativeModel('gemini-1.5-flash') # Ensured model name is standard
supabase = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_SERVICE_KEY"])

# --- 2. AUTHENTICATION ---
if 'user' not in st.session_state:
    st.markdown('<h1 style="text-align:center;">🛡️ DPIA Enterprise Portal</h1>', unsafe_allow_html=True)
    with st.container():
        c1, c2, c3 = st.columns([1, 2, 1])
        with c2:
            with st.container(border=True):
                email = st.text_input("Email address")
                password = st.text_input("Password", type="password")
                if st.button("Log In", use_container_width=True):
                    try:
                        auth = supabase.auth.sign_in_with_password({"email": email, "password": password})
                        st.session_state.user = auth.user
                        st.rerun()
                    except Exception as e:
                        st.error("Login failed")
    st.stop()

# --- 3. DASHBOARD ---
# Top Header Bar
top1, top2 = st.columns([1, 1])
top1.write(f"Logged in as: **{st.session_state.user.email}**")
if top2.button("Log Out"):
    st.session_state.clear()
    st.rerun()

# --- 4. PROJECT INPUTS ---
st.subheader("📝 New DPIA Assessment")
with st.form("dpia_form"):
    st.markdown('<div class="sub-header">I. Context & Lawfulness</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    p_name = col1.text_input("Project Name")
    p_geo = col2.text_input("Geographic Scope")
    col3, col4 = st.columns(2)
    p_basis = col3.selectbox("Lawful Basis",["Consent", "Contract", "Legal Obligation", "Legitimate Interests"])
    p_risk = col4.selectbox("Initial Risk Profile", ["Low Risk", "Medium Risk", "High Risk"])

    st.markdown('<div class="sub-header">II. Subjects & Collection</div>', unsafe_allow_html=True)
    p_subjects = st.text_input("Data Subjects")
    p_desc = st.text_area("Purpose Assessment (Processing Nature)")
    
    st.markdown('<div class="sub-header">III. Data Governance</div>', unsafe_allow_html=True)
    p_data = st.text_area("Data Points Collected")
    
    submitted = st.form_submit_button("🔍 Analyze Privacy Risks", use_container_width=True)

# --- 5. RESULTS ---
if submitted:
    with st.spinner("Analyzing with Gemini..."):
        prompt = f"Project: {p_name}. Purpose: {p_desc}. Data: {p_data}. Provide 3 risks and mitigations."
        risks = model.generate_content(prompt).text
    
    st.markdown('<div class="fb-card" style="background:white; padding:20px; border-radius:8px; border:1px solid #ddd;">', unsafe_allow_html=True)
    st.subheader("⚠️ Risk Assessment Findings")
    st.markdown(risks)
    
    # Audit Trail
    try:
        supabase.table("assessments").insert({
            "user_id": st.session_state.user.id,
            "project_name": p_name,
            "risks_output": risks,
            "created_at": str(datetime.datetime.now())
        }).execute()
    except Exception as e:
        st.error("Audit trail failed.")

    # Doc Gen
    doc = Document()
    doc.add_heading(f"DPIA Report: {p_name}", 0)
    doc.add_paragraph(risks)
    buffer = BytesIO()
    doc.save(buffer)
    
    st.download_button("📄 Download Final Word Document", buffer.getvalue(), "DPIA_Report.docx")
    st.markdown('</div>', unsafe_allow_html=True)
